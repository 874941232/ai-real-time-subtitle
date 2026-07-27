#!/usr/bin/env python3
"""Convert PRD markdown files to a single Word document."""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def read_md_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def parse_markdown(text):
    lines = text.split('\n')
    elements = []
    current_table = None
    current_paragraph = []
    
    for line in lines:
        if line.startswith('# '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('h1', line[2:]))
        
        elif line.startswith('## '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('h2', line[3:]))
        
        elif line.startswith('### '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('h3', line[4:]))
        
        elif line.startswith('#### '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('h4', line[5:]))
        
        elif line.startswith('|') and '---' in line:
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            current_table = []
        
        elif line.startswith('|') and current_table is not None:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            current_table.append(cells)
        
        elif line.startswith('```'):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('code_block', ''))
        
        elif line.startswith('- ') or line.startswith('* '):
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            elements.append(('bullet', line[2:]))
        
        elif line.strip() == '':
            if current_paragraph:
                elements.append(('paragraph', '\n'.join(current_paragraph)))
                current_paragraph = []
            if current_table:
                elements.append(('table', current_table))
                current_table = None
        
        else:
            current_paragraph.append(line)
    
    if current_paragraph:
        elements.append(('paragraph', '\n'.join(current_paragraph)))
    if current_table:
        elements.append(('table', current_table))
    
    return elements


def add_title(doc, text, level):
    if level == 1:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = '微软雅黑'
    elif level == 2:
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = '微软雅黑'
    elif level == 3:
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '微软雅黑'
    elif level == 4:
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = '微软雅黑'


def add_paragraph(doc, text):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'


def add_bullet(doc, text):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'


def add_table(doc, table_data):
    if len(table_data) < 2:
        return
    rows = len(table_data)
    cols = len(table_data[0])
    
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell
            for paragraph in table.cell(i, j).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.2
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    if i == 0:
                        run.font.bold = True


def add_code_block(doc, text):
    pass


def convert_file(doc, filepath):
    content = read_md_file(filepath)
    elements = parse_markdown(content)
    
    for elem_type, elem_content in elements:
        if elem_type == 'h1':
            add_title(doc, elem_content, 1)
        elif elem_type == 'h2':
            add_title(doc, elem_content, 2)
        elif elem_type == 'h3':
            add_title(doc, elem_content, 3)
        elif elem_type == 'h4':
            add_title(doc, elem_content, 4)
        elif elem_type == 'paragraph':
            add_paragraph(doc, elem_content)
        elif elem_type == 'bullet':
            add_bullet(doc, elem_content)
        elif elem_type == 'table':
            add_table(doc, elem_content)
        elif elem_type == 'code_block':
            add_code_block(doc, elem_content)


def main():
    prd_dir = r'c:\Users\Administrator\Desktop\0723音频\docs\prd\01_ai-real-time-subtitle'
    
    files = [
        '00_prd_ai-real-time-subtitle_overall.md',
        '01_prd_ai-real-time-subtitle_main-window.md',
        '02_prd_ai-real-time-subtitle_settings-dialog.md',
        '03_prd_ai-real-time-subtitle_subtitle-window.md',
    ]
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    for i, filename in enumerate(files):
        filepath = os.path.join(prd_dir, filename)
        if os.path.exists(filepath):
            convert_file(doc, filepath)
            if i < len(files) - 1:
                doc.add_page_break()
    
    output_path = os.path.join(prd_dir, 'AI实时字幕_PRD文档.docx')
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')


if __name__ == '__main__':
    main()
